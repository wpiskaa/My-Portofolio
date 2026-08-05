import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import re

def create_element(name):
    return OxmlElement(name)

def add_bottom_border(paragraph, color_hex="4F46E5"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="{color_hex}"/></w:pBdr>')
    pPr.append(pBdr)

# Daftar Istilah Bahasa Asing / Istilah Teknis yang Harus Di-Italic (Sesuai PUEBI/EYD)
FOREIGN_TERMS = [
    "Single Page Application", "SPA Architecture", "SPA", "Web Application", "Web App",
    "Desktop Software Engineering", "Desktop Application", "Desktop App", "Software Engineering",
    "Web Development", "UI/UX Design", "Database Management Systems", "Database",
    "Frontend Development", "Backend & Database", "Desktop & Game Dev", "Tools & UI/UX Design",
    "Live", "deploy", "di-deploy", "supplier", "real-time", "multi-spectral sensing",
    "Machine Learning", "installer", "Admin Panel", "webcam", "filter", "Pure JavaScript",
    "3D Environment Design", "Environment Design", "Publishing", "branding", "Maps Creator",
    "WebRTC", "Canvas API", "WinForms", "Import/Export", "import/export", "chart", "backup/restore",
    "POS", "Point of Sale", "Tech Stack", "Single Page Application (SPA)", "event", "Event"
]

def add_text_with_italics(paragraph, text, base_bold=False, base_italic=False, base_color=None, font_size=Pt(10.5)):
    """
    Menambahkan teks ke paragraf. Jika terdapat istilah asing dari FOREIGN_TERMS,
    istilah tersebut otomatis diformat Miring (Italic).
    """
    if not text:
        return

    # Sort foreign terms by length descending to match longest phrases first
    sorted_terms = sorted(FOREIGN_TERMS, key=len, reverse=True)
    pattern = re.compile(r'\b(' + '|'.join(re.escape(term) for term in sorted_terms) + r')\b', re.IGNORECASE)

    last_idx = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        # Normal text before match
        if start > last_idx:
            normal_part = text[last_idx:start]
            r = paragraph.add_run(normal_part)
            r.font.name = 'Calibri'
            r.font.size = font_size
            r.font.bold = base_bold
            r.font.italic = base_italic
            if base_color:
                r.font.color.rgb = base_color

        # Foreign term match (italicized)
        term_part = text[start:end]
        r_foreign = paragraph.add_run(term_part)
        r_foreign.font.name = 'Calibri'
        r_foreign.font.size = font_size
        r_foreign.font.bold = base_bold
        r_foreign.font.italic = True  # Always Italic for foreign terms
        if base_color:
            r_foreign.font.color.rgb = base_color

        last_idx = end

    # Remaining normal text
    if last_idx < len(text):
        rem_part = text[last_idx:]
        r = paragraph.add_run(rem_part)
        r.font.name = 'Calibri'
        r.font.size = font_size
        r.font.bold = base_bold
        r.font.italic = base_italic
        if base_color:
            r.font.color.rgb = base_color

def build_cv():
    doc = Document()

    # Set Margins (0.75 inch)
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Palette Warna Sesuai Desain Awal (TIDAK DIUBAH)
    PRIMARY_COLOR = RGBColor(30, 27, 75)      # Deep Indigo / Navy (#1E1B4B)
    ACCENT_COLOR = RGBColor(79, 70, 229)      # Violet / Indigo Accent (#4F46E5)
    TEXT_DARK = RGBColor(51, 65, 85)         # Slate 700 (#334155)
    TEXT_MUTED = RGBColor(100, 116, 139)     # Slate 500 (#64748B)

    # Normal Style Configuration
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = TEXT_DARK

    # -------------------------------------------------------------
    # HEADER SECTION
    # -------------------------------------------------------------
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("HAFIZ KURNIAWAN")
    run_name.font.name = 'Calibri'
    run_name.font.size = Pt(24)
    run_name.font.bold = True
    run_name.font.color.rgb = PRIMARY_COLOR

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(6)
    
    # Subtitle dengan istilah asing di-italic
    add_text_with_italics(p_title, "IT Developer & UI/UX Designer | S1 Teknologi Informasi UMY", base_bold=True, base_color=ACCENT_COLOR, font_size=Pt(11.5))

    # Contact Info Line
    p_contact = doc.add_paragraph()
    p_contact.paragraph_format.space_after = Pt(12)
    
    contacts = [
        ("Email: ", "hafizzkurniawan19@gmail.com"),
        ("LinkedIn: ", "linkedin.com/in/hafiz-kurniawan-604791370"),
        ("GitHub: ", "github.com/wpiskaa"),
        ("Instagram: ", "@hafizzkrnwn"),
        ("Lokasi: ", "Yogyakarta, Indonesia")
    ]
    
    for i, (label, val) in enumerate(contacts):
        r_lbl = p_contact.add_run(label)
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = TEXT_MUTED
        
        r_val = p_contact.add_run(val)
        r_val.font.size = Pt(9.5)
        r_val.font.color.rgb = TEXT_DARK
        
        if i < len(contacts) - 1:
            r_sep = p_contact.add_run("  •  ")
            r_sep.font.size = Pt(9.5)
            r_sep.font.color.rgb = TEXT_MUTED

    add_bottom_border(p_contact, "4F46E5")

    # Helper untuk Heading Section
    def add_section_heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        add_bottom_border(p, "CBD5E1")
        return p

    # Helper untuk menambahkan Baris Judul + Tanggal di Kanan (Strict Right Tab Stop)
    def add_item_header(paragraph, main_text, date_text, font_size=Pt(11), is_main_bold=True):
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.keep_with_next = True
        
        # Rigorous Right Tab Stop pada posisi 7.0 inchi (Sisi Kanan Margin)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)

        add_text_with_italics(paragraph, main_text, base_bold=is_main_bold, base_color=PRIMARY_COLOR, font_size=font_size)
        
        # Run Tab & Tanggal di sisi kanan
        r_tab = paragraph.add_run("\t")
        r_date = paragraph.add_run(date_text)
        r_date.font.name = 'Calibri'
        r_date.font.size = font_size
        r_date.font.bold = True
        r_date.font.color.rgb = ACCENT_COLOR

    # -------------------------------------------------------------
    # RINGKASAN PROFESIONAL (SUMMARY)
    # -------------------------------------------------------------
    add_section_heading("Ringkasan Profesional")
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_after = Pt(8)
    p_sum.paragraph_format.line_spacing = 1.15
    summary_text = (
        "Mahasiswa S1 Teknologi Informasi Universitas Muhammadiyah Yogyakarta (UMY) angkatan 2024 "
        "yang memiliki ketertarikan mendalam dan keahlian dalam Pengembangan Aplikasi Web, Aplikasi Desktop, "
        "serta Desain UI/UX. Berpengalaman memimpin divisi multimedia dalam kepanitiaan tingkat prodi/universitas, "
        "serta dipercaya oleh institusi kampus (Admisi UMY) dalam rekayasa proyek digital berskala besar. "
        "Terbiasa memadukan logika teknis yang kuat dengan estetika visual modern."
    )
    add_text_with_italics(p_sum, summary_text, base_color=TEXT_DARK)

    # -------------------------------------------------------------
    # PENDIDIKAN (EDUCATION)
    # -------------------------------------------------------------
    add_section_heading("Pendidikan")
    
    p_edu = doc.add_paragraph()
    add_item_header(p_edu, "Universitas Muhammadiyah Yogyakarta (UMY)", "2024 – Sekarang", font_size=Pt(11))

    p_deg = doc.add_paragraph()
    p_deg.paragraph_format.space_after = Pt(4)
    r_deg = p_deg.add_run("S1 Teknologi Informasi — Fakultas Teknik (NIM: 20240140024)")
    r_deg.font.italic = True
    r_deg.font.color.rgb = TEXT_MUTED
    
    p_edu_desc = doc.add_paragraph(style='List Bullet')
    p_edu_desc.paragraph_format.space_after = Pt(4)
    add_text_with_italics(p_edu_desc, "Fokus Studi: Web Application Development, Desktop Software Engineering, UI/UX Design, dan Database Management Systems.", base_color=TEXT_DARK)

    # -------------------------------------------------------------
    # KEAHLIAN TEKNIS (TECHNICAL SKILLS)
    # -------------------------------------------------------------
    add_section_heading("Keahlian & Teknologi")
    
    skills = [
        ("Frontend Development", "HTML5, CSS3, JavaScript (ES6+), React, Next.js, Tailwind CSS, TypeScript, Vue.js"),
        ("Backend & Database", "Node.js, Express, Laravel (PHP), MySQL, PostgreSQL, SQLite, REST API"),
        ("Desktop & Game Dev", "C# (WinForms), SQL Server, Lua (Roblox Studio), Canvas API (WebRTC)"),
        ("Tools & UI/UX Design", "Figma, Adobe XD, Git / GitHub, Docker, Inno Setup, Responsive Web Design")
    ]
    
    for category, items in skills:
        p_sk = doc.add_paragraph()
        p_sk.paragraph_format.space_after = Pt(3)
        
        r_bullet = p_sk.add_run("•  ")
        r_bullet.font.bold = True
        r_bullet.font.color.rgb = PRIMARY_COLOR
        
        # Category (Foreign terms italicized if applicable)
        add_text_with_italics(p_sk, f"{category}: ", base_bold=True, base_color=PRIMARY_COLOR)
        # Tech Items
        add_text_with_italics(p_sk, items, base_bold=False, base_color=TEXT_DARK)

    # -------------------------------------------------------------
    # PENGALAMAN PROYEK (KEY PROJECTS)
    # -------------------------------------------------------------
    add_section_heading("Pengalaman Proyek")

    # Format Tahun Proyek Berjalan Diubah Menjadi "Sekarang"
    projects = [
        {
            "title": "Afiyah Farmaku — Sistem Manajemen Apotek (Web SPA)",
            "date": "2025 – Sekarang",
            "tech": "JavaScript, SQLite, SPA Architecture, CSS3, Render",
            "bullets": [
                "Merekayasa aplikasi web pengelolaan apotek komprehensif berbasis Single Page Application (SPA) yang di-deploy live di Render.",
                "Mengembangkan modul Kasir (POS) interaktif real-time, manajemen stok obat modern, pencatatan pembelian dari supplier, laporan penjualan & laba-rugi, hingga manajemen hak akses karyawan dan backup/restore database."
            ]
        },
        {
            "title": "MAP UMYFIRST — UMY VERSE (Proyek Resmi Kampus UMY)",
            "date": "Juli 2025",
            "tech": "Lua, Roblox Studio, 3D Environment Design, Publishing",
            "bullets": [
                "Bertindak sebagai Maps Creator dalam proyek peta Roblox skala besar untuk inisiatif branding resmi Universitas Muhammadiyah Yogyakarta, yang ditugaskan langsung oleh Admisi UMY.",
                "Merancang lingkungan 3D interaktif kampus UMYFIRST dan menerbitkan aset game tersertifikasi resmi."
            ]
        },
        {
            "title": "HEMOSCAN & HEMO-SCAN (Aplikasi Desktop C# & Riset Medis PKM-KC)",
            "date": "2025 – Sekarang",
            "tech": "C# WinForms, SQL Server, Machine Learning, IoT, Inno Setup",
            "bullets": [
                "A8 HEMOSCAN: Mengembangkan aplikasi desktop monitoring stok darah menggunakan C# WinForms & SQL Server (Tim 2 Orang), dilengkapi fitur import/export Excel, chart statistik, dan installer Inno Setup.",
                "HEMO-SCAN (PKM-KC UMY): Meriset prototipe alat medis portabel pendeteksi golongan darah (ABO & Rhesus) tanpa reagen berbasis 11-channel multi-spectral sensing & Machine Learning."
            ]
        },
        {
            "title": "Sistem Manajemen Kepanitiaan IT Specta 2026",
            "date": "2026 – Sekarang",
            "tech": "Web Application, HTML/CSS/JS, Admin Panel",
            "bullets": [
                "Membangun dan mengimplementasikan sistem web manajemen internal untuk koordinasi kepanitiaan dan operasional acara besar IT Specta 2026 Prodi TI UMY."
            ]
        },
        {
            "title": "Browser Game & Photobooth In-Browser (MyWorkSpace Modules)",
            "date": "2026 – Sekarang",
            "tech": "JavaScript, Canvas API, WebRTC, HTML5",
            "bullets": [
                "Browser Game (Chill): Mengembangkan mini-game interaktif 2D playable langsung di browser menggunakan Pure JavaScript & Canvas API.",
                "Photobooth In-Browser: Membangun modul photobooth interaktif dengan integrasi webcam (WebRTC), filter foto real-time, dan fitur ekspor gambar langsung."
            ]
        }
    ]

    for proj in projects:
        p_p = doc.add_paragraph()
        add_item_header(p_p, proj["title"], proj["date"], font_size=Pt(11))

        p_tc = doc.add_paragraph()
        p_tc.paragraph_format.space_after = Pt(3)
        r_tc_lbl = p_tc.add_run("Teknologi: ")
        r_tc_lbl.font.bold = True
        r_tc_lbl.font.size = Pt(9.5)
        r_tc_lbl.font.color.rgb = TEXT_MUTED
        add_text_with_italics(p_tc, proj["tech"], base_color=ACCENT_COLOR, font_size=Pt(9.5))

        for bullet in proj["bullets"]:
            p_b = doc.add_paragraph(style='List Bullet')
            p_b.paragraph_format.space_after = Pt(3)
            p_b.paragraph_format.line_spacing = 1.15
            add_text_with_italics(p_b, bullet, base_color=TEXT_DARK)

    # -------------------------------------------------------------
    # PENGALAMAN ORGANISASI & KEPANITIAAN
    # -------------------------------------------------------------
    add_section_heading("Pengalaman Organisasi & Kepanitiaan")

    orgs = [
        {
            "role": "Koordinator Divisi Multimedia",
            "event": "IT Specta 2026 — Himpunan / Prodi Teknologi Informasi UMY",
            "date": "2026 – Sekarang",
            "desc": "Memimpin dan memandu tim multimedia dalam produksi materi promosi visual, aset dekorasi digital, media publikasi, serta pengembangan sistem internal kepanitiaan."
        },
        {
            "role": "Anggota Panitia",
            "event": "IT Specta 2025 — Program Studi Teknologi Informasi UMY",
            "date": "2025",
            "desc": "Mengelola persiapan teknis dan operasional lapangan selama rangkaian kegiatan IT Specta 2025."
        },
        {
            "role": "Anggota Panitia",
            "event": "IT Phoria Fest 2024 — Program Studi Teknologi Informasi UMY",
            "date": "2024",
            "desc": "Mengonsolidasikan logistik dan dukungan teknis untuk kesuksesan event tahunan prodi."
        },
        {
            "role": "Panitia & Peserta",
            "event": "Seminar Nasional: Manusia vs Mesin",
            "date": "2025",
            "desc": "Berpartisipasi aktif dalam penyelenggaraan seminar nasional yang membahas kecerdasan buatan dan dampaknya terhadap tenaga kerja manusia."
        }
    ]

    for org in orgs:
        p_o = doc.add_paragraph()
        full_org_title = f"{org['role']} — {org['event']}"
        add_item_header(p_o, full_org_title, org["date"], font_size=Pt(10.5))

        p_od = doc.add_paragraph(style='List Bullet')
        p_od.paragraph_format.space_after = Pt(4)
        add_text_with_italics(p_od, org["desc"], base_color=TEXT_DARK)

    # -------------------------------------------------------------
    # PRESTASI & REKOGNISI (RECOGNITIONS & COMPETITIONS)
    # -------------------------------------------------------------
    add_section_heading("Prestasi & Rekognisi")

    achievements = [
        ("Kompetisi UI/UX Arkavidia (UXVIDIA ITB)", "Februari 2025", "Lolos babak penyisihan kompetisi UI/UX tingkat Nasional yang diselenggarakan oleh HMIF Institut Teknologi Bandung (ITB)."),
        ("Find IT UGM (Esports - VALORANT)", "2025", "Berpartisipasi dalam babak penyisihan ajang kompetisi teknologi dan minat bakat tahunan Universitas Gadjah Mada (UGM).")
    ]

    for name, date, desc in achievements:
        p_a = doc.add_paragraph()
        add_item_header(p_a, name, date, font_size=Pt(10.5))

        p_ad_desc = doc.add_paragraph(style='List Bullet')
        p_ad_desc.paragraph_format.space_after = Pt(4)
        add_text_with_italics(p_ad_desc, desc, base_color=TEXT_DARK)

    # Save document
    output_path = "D:\\KULIAH\\CV N PORTOFOLIO\\CV_Hafiz_Kurniawan.docx"
    doc.save(output_path)
    print(f"CV updated successfully to {output_path}")

if __name__ == "__main__":
    build_cv()
